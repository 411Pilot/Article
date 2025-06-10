import streamlit as st
import google.generativeai as genai
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import textstat
import re
import markdown
import docx
import os
from datetime import datetime

# Configure Gemini API
genai.configure(api_key="AIzaSyBo4qC13bRaJDJt1_NcUeEX2FfwN6ghK-U")
model = genai.GenerativeModel("gemini-1.5-flash")

# Streamlit App Configuration
st.set_page_config(page_title="LinkedIn Blog Generator", layout="wide")
st.title("AI-Powered LinkedIn Blog & Article Generator")
st.markdown("Create professional blogs and LinkedIn posts with AI-powered content generation.")

# Initialize session state for content storage
if "generated_content" not in st.session_state:
    st.session_state.generated_content = {}
if "blog_title" not in st.session_state:
    st.session_state.blog_title = ""
if "blog_content" not in st.session_state:
    st.session_state.blog_content = ""

# Sidebar for Input
with st.sidebar:
    st.header("Content Input")
    user_input = st.text_area("Enter your blog topic or key points:", height=150)
    article_template = st.selectbox(
        "Select Article Template",
        ["How-to", "Listicle", "Opinion", "Case Study"]
    )
    tone = st.selectbox(
        "Select Tone",
        ["Professional", "Friendly", "Casual", "Witty", "Technical"]
    )
    audience = st.selectbox(
        "Target Audience",
        ["HRs", "Tech Founders", "Marketers", "Freshers", "General"]
    )
    primary_keyword = st.text_input("Primary SEO Keyword (optional):")
    generate_button = st.button("Generate Content")
    regenerate_button = st.button("Regenerate with New Tone/Template")

# Function to generate content using Gemini API
def generate_content(prompt):
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating content: {str(e)}")
        return ""

# Function to generate blog structure
def generate_blog_structure(template, topic, tone, audience):
    prompt = f"""
    Create a blog outline for a {template.lower()} article on '{topic}'.
    Target audience: {audience}.
    Tone: {tone}.
    Include:
    - Blog title
    - Introduction
    - 3-5 main sections with subheadings
    - Conclusion
    """
    return generate_content(prompt)

# Function to generate LinkedIn post summary
def generate_linkedin_summary(content):
    prompt = f"Generate a concise LinkedIn post summary (TL;DR style, max 150 words) from the following blog content:\n{content}"
    return generate_content(prompt)

# Function to generate quote cards
def generate_quote_cards(content):
    prompt = f"Extract 3 tweetable quotes (max 280 characters each) from the following blog content:\n{content}"
    return generate_content(prompt)

# ✅ Fixed Function Name
def suggest_seo_metadata(title, keyword):
    prompt = f"Suggest an SEO-friendly title (up to 60 characters) and meta description (up to 160 characters) for a blog with the keyword '{keyword}'."
    return generate_content(prompt)

# Function to check readability
def check_readability(text):
    score = textstat.flesch_kincaid_grade(text)
    return f"Flesch-Kincaid Grade Level: {score:.2f}"

# Function to create PDF
def create_pdf(title, content):
    pdf_file = f"{title}.pdf"
    doc = SimpleDocTemplate(pdf_file, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))
    for line in content.split("\n"):
        story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 12))
    doc.build(story)
    return pdf_file

# Function to create DOC
def create_doc(title, content):
    doc = docx.Document()
    doc.add_heading(title, 0)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc_file = f"{title}.docx"
    doc.save(doc_file)
    return doc_file

# Function to convert to HTML
def convert_to_html(content):
    html_content = markdown.markdown(content)
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Blog Post</title>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

# Main Content Generation
if generate_button or regenerate_button:
    if user_input:
        # Generate blog structure
        blog_content = generate_blog_structure(article_template, user_input, tone, audience)
        st.session_state.blog_content = blog_content
        st.session_state.blog_title = re.search(r"#+ (.*?)\n", blog_content).group(1) if re.search(r"#+ (.*?)\n", blog_content) else "Generated Blog"

        # Display generated content
        st.subheader("Generated Blog Content")
        st.markdown(blog_content)

        # LinkedIn Post Summary
        st.subheader("LinkedIn Post Summary")
        linkedin_summary = generate_linkedin_summary(blog_content)
        st.markdown(linkedin_summary)
        st.session_state.generated_content["linkedin_summary"] = linkedin_summary

        # Quote Cards
        st.subheader("Tweetable Quote Cards")
        quotes = generate_quote_cards(blog_content)
        st.markdown(quotes)
        st.session_state.generated_content["quotes"] = quotes

        # SEO Suggestions
        if primary_keyword:
            st.subheader("SEO Suggestions")
            seo_suggestions = suggest_seo_metadata(st.session_state.blog_title, primary_keyword)
            st.markdown(seo_suggestions)
            st.session_state.generated_content["seo"] = seo_suggestions

        # Readability Check
        st.subheader("Readability Analysis")
        readability = check_readability(blog_content)
        st.write(readability)

        # Image Suggestions (Placeholder)
        st.subheader("Image Suggestions")
        st.write("Suggested royalty-free image sources: Unsplash, Pexels (use keyword: " + (primary_keyword or user_input) + ")")
        st.session_state.generated_content["image_suggestion"] = "Royalty-free images from Unsplash or Pexels"

        # Plagiarism Check (Placeholder)
        st.subheader("Plagiarism Check")
        st.write("Plagiarism checking requires integration with external APIs (e.g., Copyscape).")

# Export Options
st.header("Export Options")
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("Download as PDF"):
        if st.session_state.blog_content:
            pdf_file = create_pdf(st.session_state.blog_title, st.session_state.blog_content)
            with open(pdf_file, "rb") as f:
                st.download_button(
                    label="Download PDF",
                    data=f,
                    file_name=pdf_file,
                    mime="application/pdf"
                )

with col2:
    if st.button("Download as DOC"):
        if st.session_state.blog_content:
            doc_file = create_doc(st.session_state.blog_title, st.session_state.blog_content)
            with open(doc_file, "rb") as f:
                st.download_button(
                    label="Download DOC",
                    data=f,
                    file_name=doc_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

with col3:
    if st.button("Export as HTML"):
        if st.session_state.blog_content:
            html_content = convert_to_html(st.session_state.blog_content)
            st.download_button(
                label="Download HTML",
                data=html_content,
                file_name="blog.html",
                mime="text/html"
            )

# Scheduling Placeholder
st.subheader("Schedule Post")
st.write("Scheduling requires LinkedIn API integration. Save the content and use LinkedIn's native scheduler.")
